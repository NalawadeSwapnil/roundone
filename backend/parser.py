"""
parser.py — CV Parser for Roundone
-----------------------------------
Extracts and cleans text from a PDF or DOCX CV.

Libraries used:
  - pdfplumber  : PDF parsing — handles multi-column layouts reliably
  - python-docx : DOCX parsing — reads paragraphs and tables from Word docs
"""

import re
import os
import pdfplumber
from docx import Document


def _clean(raw_text: str) -> str:
    """
    Shared cleaning pipeline applied to text from any file format.

    Steps:
      1. Normalise line endings (Windows \r\n → \n)
      2. Collapse 3+ consecutive blank lines into one blank line
      3. Strip leading/trailing whitespace from every line
      4. Collapse multiple spaces/tabs within a line into a single space
      5. Final strip of the whole string
    """
    clean = raw_text.replace("\r\n", "\n").replace("\r", "\n")
    clean = re.sub(r"\n{3,}", "\n\n", clean)
    clean = "\n".join(line.strip() for line in clean.splitlines())
    clean = re.sub(r"[ \t]{2,}", " ", clean)
    return clean.strip()


def _parse_pdf(file_path: str) -> str:
    """Extract text from every page of a PDF file."""
    extracted_pages = []

    with pdfplumber.open(file_path) as pdf:
        for page_number, page in enumerate(pdf.pages, start=1):
            text = page.extract_text()
            if text:
                extracted_pages.append(text)
            else:
                print(f"  [Warning] Page {page_number} returned no text — skipping.")

    if not extracted_pages:
        raise ValueError(
            "No extractable text found in the PDF. "
            "It may be a scanned image. Please use a text-based PDF."
        )

    return "\n".join(extracted_pages)


def _parse_docx(file_path: str) -> str:
    """
    Extract text from a DOCX file.

    Reads two sources:
      - Paragraphs : body text, headings, bullet points
      - Tables     : skills grids, education/experience tables

    Why both? CVs often use invisible tables for layout (e.g. two-column
    skills sections). Extracting table cells separately ensures that content
    is not silently dropped.
    """
    doc = Document(file_path)
    lines = []

    # 1. Extract all paragraph text
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            lines.append(text)

    # 2. Extract text from all table cells
    for table in doc.tables:
        for row in table.rows:
            row_cells = []
            for cell in row.cells:
                cell_text = cell.text.strip()
                if cell_text:
                    row_cells.append(cell_text)
            if row_cells:
                # Join cells in the same row with a separator so they read naturally
                lines.append("  |  ".join(row_cells))

    if not lines:
        raise ValueError(
            "No extractable text found in the DOCX file. "
            "The document may be empty or contain only images."
        )

    return "\n".join(lines)


def parse_cv(file_path: str) -> str:
    """
    Extract and clean all text from a CV file.

    Supports:
      - .pdf  (via pdfplumber)
      - .docx (via python-docx)

    Args:
        file_path: Absolute or relative path to the CV file.

    Returns:
        A single clean string containing all text from the CV.

    Raises:
        FileNotFoundError : If the file does not exist.
        ValueError        : If the format is unsupported or no text is found.
    """
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"File not found: '{file_path}'")

    extension = os.path.splitext(file_path)[1].lower()

    if extension == ".pdf":
        raw_text = _parse_pdf(file_path)
    elif extension == ".docx":
        raw_text = _parse_docx(file_path)
    else:
        raise ValueError(
            f"Unsupported file format: '{extension}'. "
            "Please upload a .pdf or .docx file."
        )

    return _clean(raw_text)


# ---------------------------------------------------------------------------
# Quick self-test — run this file directly to verify the parser works:
#   python parser.py your_cv.pdf
#   python parser.py your_cv.docx
# ---------------------------------------------------------------------------
if __name__ == "__main__":
    import sys

    test_path = sys.argv[1] if len(sys.argv) > 1 else "sample_cv.pdf"

    if not os.path.exists(test_path):
        print(f"[Test] File not found: '{test_path}'")
        print("       Pass a file path as an argument:")
        print("         python parser.py your_cv.pdf")
        print("         python parser.py your_cv.docx")
        sys.exit(1)

    print(f"[Test] Parsing: {test_path}")
    print("-" * 60)

    result = parse_cv(test_path)

    print(result[:1000])  # Preview first 1000 characters
    print("-" * 60)
    print(f"[Test] Total characters extracted : {len(result)}")
    print(f"[Test] Total lines extracted      : {result.count(chr(10)) + 1}")
    print("[Test] ✓ Parser working correctly.")
